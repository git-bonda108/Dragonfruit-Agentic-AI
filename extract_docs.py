#!/usr/bin/env python3
"""Extract text from xlsx and docx for analysis."""
import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document
import json
import os

BASE = os.path.dirname(os.path.abspath(__file__))

def extract_xlsx(path):
    wb = openpyxl.load_workbook(path, data_only=True, read_only=False)
    out = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        out.append(f"\n=== Sheet: {sheet_name} ===\n")
        for row in ws.iter_rows(values_only=True):
            row_vals = [str(c) if c is not None else "" for c in row]
            if any(row_vals):
                out.append("\t".join(row_vals))
    return "\n".join(out)

def extract_docx(path):
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs) + "\n\n" + "\n".join(
        " | ".join(c.text for c in row.cells) for table in doc.tables for row in table.rows
    )

files = [
    ("DFM Current Pipeline Tool Stack.xlsx", extract_xlsx),
    ("DFM_Current_vs_Dream_Time_Allocations.xlsx", extract_xlsx),
    ("DFM Current Pipeline.xlsx", extract_xlsx),
    ("DFM_AI_Tool_Wishlist_Consultant_Brief.docx", extract_docx),
]

for fname, fn in files:
    p = os.path.join(BASE, fname)
    if os.path.exists(p):
        try:
            text = fn(p)
            out_path = os.path.join(BASE, "extracted_" + os.path.splitext(fname)[0] + ".txt")
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(text)
            print(f"Wrote {out_path}")
        except Exception as e:
            print(f"Error {fname}: {e}")
    else:
        print(f"Missing: {p}")
